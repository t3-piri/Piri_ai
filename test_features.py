"""4 yeni ozelligin (kaynak/versiyon takibi, guven seviyesi, kanit-yetersizse-
yonlendirme) 3 senaryo ile testi. Gercek admin_panel.py'nin input() akisina
girmeden, ayni fonksiyonlari dogrudan cagirarak calisir."""

from pathlib import Path

from docx import Document

from document_registry import make_document_id, register_new_version
from local_ingest import get_collection, get_tokenizer, ingest_records, records_for_file
import local_rag_answer
from local_rag_answer import answer_question, search_competition
from qa_log import unanswered_questions

ROOT = Path("Piri-veriler")
TEST_COMPETITION = "_Test Versiyonlama"


def ok(label, condition):
    print(f"  [{'PASS' if condition else 'FAIL'}] {label}")
    return condition


def test_a_high_confidence():
    print("\n=== SENARYO A: Yuksek guvenle cevaplanabilen soru ===")
    result = answer_question(
        "Roket yarışmasında takımda en az en fazla kaç kişi olabilir?", "Roket Yarışması"
    )
    print("Yanit:", result["answer"][:300])
    all_ok = True
    all_ok &= ok("status == answered", result["status"] == "answered")
    all_ok &= ok("yanitta 'Kaynak: [' var", "Kaynak: [" in result["answer"])
    all_ok &= ok("yanitta 'Güven seviyesi:' var", "Güven seviyesi:" in result["answer"])
    all_ok &= ok("guven seviyesi Orta/Yuksek", result["confidence"] in ("Orta güven", "Yüksek güven"))
    return all_ok


def test_b_low_confidence():
    print("\n=== SENARYO B: Belirsiz/alakasiz soru ===")
    question = "Roket yarışması bağlamında marsta koloni kurma stratejisi nasıl olmalı ve evlilik törenleri nasıl düzenlenir?"
    result = answer_question(question, "Roket Yarışması")
    print("Yanit:", result["answer"][:300])
    all_ok = True
    all_ok &= ok("status == low_confidence", result["status"] == "low_confidence")
    unanswered = [e for e in unanswered_questions() if e["question"] == question]
    all_ok &= ok("unanswered_questions'a kaydedildi", len(unanswered) >= 1)
    return all_ok


def _make_test_docx(path, sentence):
    doc = Document()
    doc.add_heading("Test Belgesi", level=1)
    doc.add_paragraph(sentence)
    doc.add_paragraph(sentence)
    doc.add_paragraph(sentence)
    doc.save(path)


def test_c_versioning():
    print("\n=== SENARYO C: Ayni belgenin guncellenmis versiyonu ===")
    test_dir = ROOT / TEST_COMPETITION
    test_dir.mkdir(parents=True, exist_ok=True)
    file_name = "versiyon_testi.docx"
    file_path = test_dir / file_name
    rel_path = str(file_path.relative_to(ROOT))
    document_id = make_document_id(TEST_COMPETITION, file_name)

    collection = get_collection()
    tokenizer = get_tokenizer()

    # v1
    _make_test_docx(file_path, "Gizli kod kelimesi ALFA111 burada gecer.")
    doc_id, v1 = register_new_version(TEST_COMPETITION, file_name, rel_path)
    v1_records = records_for_file(file_path, TEST_COMPETITION, tokenizer, version=v1, status="active", document_id=doc_id)
    ingest_records(v1_records, collection)
    local_rag_answer._bm25_cache.pop(TEST_COMPETITION, None)

    hits_v1, score_v1 = search_competition("ALFA111 gizli kod kelimesi", TEST_COMPETITION, top_k=3)
    v1_found = any("ALFA111" in h["text"] for h in hits_v1)

    # v2: aynı dosya adi, farkli icerik -> eski versiyon inactive olmali
    old = collection.get(
        where={"$and": [{"source_path": rel_path}, {"status": "active"}]}, include=["metadatas"]
    )
    collection.update(ids=old["ids"], metadatas=[{**m, "status": "inactive"} for m in old["metadatas"]])

    _make_test_docx(file_path, "Gizli kod kelimesi BETA222 burada gecer.")
    doc_id, v2 = register_new_version(TEST_COMPETITION, file_name, rel_path)
    v2_records = records_for_file(file_path, TEST_COMPETITION, tokenizer, version=v2, status="active", document_id=doc_id)
    ingest_records(v2_records, collection)
    local_rag_answer._bm25_cache.pop(TEST_COMPETITION, None)

    hits_after, _ = search_competition("ALFA111 gizli kod kelimesi", TEST_COMPETITION, top_k=5)
    old_still_active = any("ALFA111" in h["text"] for h in hits_after)

    hits_v2, _ = search_competition("BETA222 gizli kod kelimesi", TEST_COMPETITION, top_k=5)
    v2_found = any("BETA222" in h["text"] for h in hits_v2)

    all_ok = True
    all_ok &= ok("v1 ilk yuklemede aramada bulunuyordu", v1_found)
    all_ok &= ok("v2 yuklendikten sonra v1 ARTIK aramada YOK", not old_still_active)
    all_ok &= ok("v2 aramada bulunuyor", v2_found)
    all_ok &= ok("yeni versiyon numarasi v1'den buyuk", v2 > v1)
    return all_ok


def cleanup_test_artifacts():
    import shutil
    import sqlite3

    collection = get_collection()
    collection.delete(where={"competition": TEST_COMPETITION})
    local_rag_answer._bm25_cache.pop(TEST_COMPETITION, None)

    conn = sqlite3.connect("documents.db")
    conn.execute("DELETE FROM documents WHERE competition = ?", (TEST_COMPETITION,))
    conn.commit()
    conn.close()

    test_dir = ROOT / TEST_COMPETITION
    if test_dir.exists():
        shutil.rmtree(test_dir)

    # local_ingest_progress.txt'te bu test'in id'leri "islenmis" olarak kalirsa,
    # bir sonraki calistirmada ingest_records() Chroma'dan silinen bu chunk'lari
    # "zaten islenmis" sanip atlar ve v1 hic eklenmeden test yanlislikla basarisiz
    # olur. Checkpoint'i Chroma'nin gercek durumuyla senkronize ederek onlenir.
    checkpoint_file = Path("local_ingest_progress.txt")
    if checkpoint_file.exists():
        all_ids = set(collection.get(include=[])["ids"])
        remaining = [i for i in checkpoint_file.read_text(encoding="utf-8").splitlines() if i in all_ids]
        checkpoint_file.write_text("\n".join(remaining) + ("\n" if remaining else ""), encoding="utf-8")

    print(f"\nTest verileri temizlendi ({TEST_COMPETITION}).")


if __name__ == "__main__":
    results = {
        "A": test_a_high_confidence(),
        "B": test_b_low_confidence(),
        "C": test_c_versioning(),
    }
    print("\n===== SONUC =====")
    for k, v in results.items():
        print(f"Senaryo {k}: {'BASARILI' if v else 'BASARISIZ'}")

    cleanup_test_artifacts()
